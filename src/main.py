import sys
import os
import warnings
import time
import pickle
from pathlib import Path
from typing import List
from concurrent.futures import ThreadPoolExecutor

# ==============================
# ⚙️ Configurações principais
# ==============================
DEFAULT_DOCUMENTS_FOLDER = "documentos"
FAISS_INDEX_DIR = "./storage"
CACHE_FILE = "./storage/rag_cache.pkl"
EMBEDDING_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"
OLLAMA_MODEL_NAME = "llama3.2:3b"

CHUNK_SIZE = 512
CHUNK_OVERLAP = 50
SIMILARITY_TOP_K = 3
MAX_TOKENS = 512
NUM_CTX = 2048

SYSTEM_PROMPT = """Responda em português, de forma direta e objetiva.
Liste todos os itens mencionados no contexto.
Seja breve e preciso.
"""

EXIT_COMMANDS = ["sair", "exit", "quit"]

# ==============================
# 🚫 Supressão de avisos
# ==============================
warnings.filterwarnings("ignore")
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"

# ==============================
# 🧠 Funções utilitárias
# ==============================
def check_ollama_running():
    import requests
    try:
        resp = requests.get("http://localhost:11434/api/tags", timeout=5)
        return resp.status_code == 200
    except Exception:
        return False


def check_gpu_available():
    try:
        import torch
        if torch.cuda.is_available():
            print(f"🎮 GPU detectada: {torch.cuda.get_device_name(0)}")
            return True
        else:
            print("💻 Usando CPU")
            return False
    except Exception:
        print("💻 Usando CPU")
        return False


# ==============================
# 📚 Carregamento de documentos (paralelizado)
# ==============================
def load_file(f):
    try:
        suffix = f.suffix.lower()
        if suffix == ".txt":
            return f.read_text(encoding="utf-8")
        elif suffix == ".docx":
            from docx import Document as DocxDocument
            docx = DocxDocument(str(f))
            return "\n".join(p.text for p in docx.paragraphs if p.text.strip())
        elif suffix == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(str(f))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        return None
    except Exception:
        return None


def load_documents(folder_path: str):
    from llama_index.core import Document

    folder = Path(folder_path)
    files = list(folder.glob("*"))
    if not files:
        raise ValueError(f"A pasta '{folder_path}' está vazia.")

    with ThreadPoolExecutor(max_workers=4) as ex:
        contents = list(ex.map(load_file, files))

    docs = [
        Document(text=c, doc_id=f.name, metadata={"source": str(f)})
        for f, c in zip(files, contents)
        if c and c.strip()
    ]

    print(f"✅ {len(docs)} documentos carregados.")
    return docs


# ==============================
# 💾 Cache local de embeddings e índice
# ==============================
def save_cache(index, embed_model, llm):
    Path(FAISS_INDEX_DIR).mkdir(exist_ok=True)
    with open(CACHE_FILE, "wb") as f:
        pickle.dump((index, embed_model, llm), f)


def load_cache():
    if os.path.exists(CACHE_FILE):
        with open(CACHE_FILE, "rb") as f:
            return pickle.load(f)
    return None, None, None


# ==============================
# ⚡ Criação do índice vetorial
# ==============================
def create_faiss_vector_store(embedding_dim=384):
    import faiss
    from llama_index.vector_stores.faiss import FaissVectorStore

    quantizer = faiss.IndexFlatL2(embedding_dim)
    index = faiss.IndexIVFFlat(quantizer, embedding_dim, 100)
    return FaissVectorStore(faiss_index=index)


# ==============================
# 🧩 Configuração do RAG
# ==============================
def setup_rag_system(documents_folder=DEFAULT_DOCUMENTS_FOLDER):
    from llama_index.core import (
        VectorStoreIndex,
        Settings,
        StorageContext,
    )
    from llama_index.embeddings.huggingface import HuggingFaceEmbedding
    from llama_index.llms.ollama import Ollama

    print("📦 Inicializando RAG...")

    # 1️⃣ Tenta carregar cache
    index, embed_model, llm = load_cache()
    if index and embed_model and llm:
        print("✅ Cache carregado com sucesso.")
        return index.as_query_engine(
            similarity_top_k=SIMILARITY_TOP_K,
            response_mode="compact",
            streaming=False,
        )

    # 2️⃣ Carrega documentos
    documents = load_documents(documents_folder)

    # 3️⃣ Embeddings
    print("🔧 Carregando modelo de embeddings...")
    embed_model = HuggingFaceEmbedding(
        model_name=EMBEDDING_MODEL_NAME,
        cache_folder="./.cache/embeddings",
        device="cuda" if check_gpu_available() else "cpu",
    )

    # 4️⃣ LLM (Ollama local)
    print("🔧 Configurando LLM Ollama...")
    llm = Ollama(
        model=OLLAMA_MODEL_NAME,
        request_timeout=60.0,
        temperature=0.0,
        system_prompt=SYSTEM_PROMPT,
        additional_kwargs={
            "num_predict": MAX_TOKENS,
            "num_ctx": NUM_CTX,
            "num_thread": os.cpu_count(),
        },
    )

    # 5️⃣ Configuração global
    Settings.llm = llm
    Settings.embed_model = embed_model
    Settings.chunk_size = CHUNK_SIZE
    Settings.chunk_overlap = CHUNK_OVERLAP

    # 6️⃣ Cria índice FAISS
    print("🔧 Construindo FAISS...")
    vector_store = create_faiss_vector_store(embedding_dim=384)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    # 7️⃣ Cria índice vetorial
    print("⚙️ Criando índice...")
    index = VectorStoreIndex.from_documents(
        documents,
        storage_context=storage_context,
        show_progress=True,
    )

    index.storage_context.persist(persist_dir=FAISS_INDEX_DIR)
    save_cache(index, embed_model, llm)

    print("✅ Sistema RAG configurado e cache salvo.")
    return index.as_query_engine(
        similarity_top_k=SIMILARITY_TOP_K,
        response_mode="compact",
        streaming=False,
    )


# ==============================
# 💬 Loop principal do chatbot
# ==============================
def run_chat_loop(query_engine):
    print("\n" + "=" * 40)
    print("🤖 Chatbot RAG — digite sua pergunta")
    print("Digite 'sair' para encerrar.")
    print("=" * 40 + "\n")

    while True:
        try:
            pergunta = input("👤 Você: ").strip()
            if not pergunta:
                continue
            if pergunta.lower() in EXIT_COMMANDS:
                print("👋 Encerrando.")
                break

            print("🔍 Processando...")
            start = time.time()

            resposta = query_engine.query(pergunta)

            print(f"\n🤖 Bot: {resposta}")
            print(f"⏱️ Tempo: {time.time() - start:.2f}s\n")

        except KeyboardInterrupt:
            print("\n👋 Interrompido.")
            break
        except Exception as e:
            print(f"❌ Erro: {e}")


# ==============================
# 🚀 Main
# ==============================
def main():
    print("=" * 60)
    print("🤖 RAG Chatbot - Inicializando...")
    print("=" * 60 + "\n")

    print(f"📋 Verificando modelo: {OLLAMA_MODEL_NAME}")
    if not check_ollama_running():
        print("❌ Ollama não está rodando. Rode: `ollama serve`")
        sys.exit(1)

    # Configura e roda
    try:
        query_engine = setup_rag_system()
        run_chat_loop(query_engine)
    except Exception as e:
        print(f"❌ Erro crítico: {e}")
        raise


if __name__ == "__main__":
    main()
